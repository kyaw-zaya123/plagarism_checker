#main.py
import os
import docx
import pdfplumber
from html2text import html2text
from flask_bcrypt import Bcrypt
from functools import wraps
from werkzeug.security import generate_password_hash
from flask_login import LoginManager, UserMixin, login_user, login_required, logout_user, current_user
from flask import Flask, request, render_template, redirect, flash, url_for, jsonify
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import nltk
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize
from nltk.stem import WordNetLemmatizer
import mysql.connector
from mysql.connector import Error
import datetime
import json
from flasgger import Swagger, swag_from

nltk.download('punkt')
nltk.download('wordnet')
nltk.download('stopwords')

app = Flask(__name__, template_folder='templates')
app.secret_key = '9393c60074ab53a68d814334382cff7f'
bcrypt = Bcrypt(app)
Swagger(app)

from flask_login import UserMixin
login_manager = LoginManager(app)
login_manager.login_view = 'login'

class User(UserMixin):
    def __init__(self, id, username, password, is_admin=False):
        self.id = id
        self.username = username
        self.password = password
        self.is_admin = is_admin

lemmatizer = WordNetLemmatizer()
stop_words = set(stopwords.words('russian'))

UPLOAD_FOLDER = 'project-file'

DB_CONFIG = {
    'host': '127.0.0.1',
    'user': 'root',
    'password': 'Kyaw550550#',
    'database': 'my_new_database'
}

RECORDS_PER_PAGE = 10

def create_database_connection():
    """Create a database connection to MySQL"""
    try:
        connection = mysql.connector.connect(**DB_CONFIG)
        if connection.is_connected():
            return connection
    except Error as e:
        print(f"Error while connecting to MySQL: {e}")
    return None

@login_manager.user_loader
def load_user(user_id):
    connection = create_database_connection()
    if connection:
        try:
            cursor = connection.cursor(dictionary=True)
            cursor.execute("SELECT * FROM users WHERE id = %s", (user_id,))
            user = cursor.fetchone()
            if user:
                return User(id=user['id'], username=user['username'], password=user['password'], is_admin=user['is_admin'])
        except Error as e:
            print(f"Error loading user: {e}")
        finally:
            if connection.is_connected():
                cursor.close()
                connection.close()
    return None

def alter_comparisons_table():
    """Alter the comparisons table to add AUTO_INCREMENT to the id field"""
    connection = create_database_connection()
    if connection:
        try:
            cursor = connection.cursor()
            cursor.execute("""
                ALTER TABLE comparisons MODIFY COLUMN id INT AUTO_INCREMENT
            """)
            connection.commit()
            print("Comparisons table altered successfully.")
        except Error as e:
            print(f"Error altering comparisons table: {e}")
        finally:
            if connection.is_connected():
                cursor.close()
                connection.close()

def ensure_upload_folder_exists():
    """Ensure the upload folder exists."""
    if not os.path.exists(UPLOAD_FOLDER):
        os.makedirs(UPLOAD_FOLDER)

def preprocess_text(text):
    """Preprocess text by tokenizing, removing stop words, and lemmatizing."""
    tokens = word_tokenize(text.lower())
    tokens = [lemmatizer.lemmatize(token) for token in tokens if token not in stop_words and token.isalnum()]
    return ' '.join(tokens)

def convert_file(file_path, skip_first_page=True):
    """Convert file to text based on its type and preprocess the text, optionally skipping the first page."""
    file_extension = file_path.rsplit('.', 1)[-1].lower()
    
    if file_extension == 'docx':
        doc = docx.Document(file_path)
        if skip_first_page:
            second_page_start = 0
            for i, para in enumerate(doc.paragraphs):
                if para.runs and para.runs[0].element.xpath('.//w:br[@w:type="page"]'):
                    second_page_start = i + 1
                    break
            text = '\n'.join([para.text for para in doc.paragraphs[second_page_start:]])
        else:
            text = '\n'.join([para.text for para in doc.paragraphs])
    elif file_extension == 'pdf':
        with pdfplumber.open(file_path) as pdf:
            pages = pdf.pages[1:] if skip_first_page else pdf.pages
            text = '\n'.join([page.extract_text() for page in pages if page.extract_text()])
    elif file_extension == 'html':
        with open(file_path, 'r', encoding='utf-8') as html_file:
            html_content = html_file.read()
            text = html2text(html_content)
            if skip_first_page:
                text = text[1500:]
    else:
        with open(file_path, 'r', encoding='utf-8') as text_file:
            lines = text_file.readlines()
            if skip_first_page:
                text = ''.join(lines[50:])
            else:
                text = ''.join(lines)
    return preprocess_text(text)

def save_file_to_db(filename, content, user_id):
    """Save file information to the database"""
    connection = create_database_connection()
    if connection:
        try:
            cursor = connection.cursor()
            query = "INSERT INTO files (filename, content, user_id) VALUES (%s, %s, %s)"
            cursor.execute(query, (filename, content, user_id))
            connection.commit()
            return cursor.lastrowid
        except Error as e:
            print(f"Error saving file to database: {e}")
            connection.rollback()
        finally:
            if connection.is_connected():
                cursor.close()
                connection.close()
    return None

def save_comparison_to_db(file1_id, file2_id, similarity, user_id, highlighted_content1, highlighted_content2):
    """Save comparison results to the database"""
    connection = create_database_connection()
    if connection:
        try:
            cursor = connection.cursor()
            query = """
            INSERT INTO comparisons 
            (file1_id, file2_id, similarity, user_id, highlighted_content1, highlighted_content2) 
            VALUES (%s, %s, %s, %s, %s, %s)
            """
            cursor.execute(query, (file1_id, file2_id, float(similarity), user_id, highlighted_content1, highlighted_content2))
            connection.commit()
            print(f"Comparison saved: file1_id={file1_id}, file2_id={file2_id}, similarity={similarity}, user_id={user_id}")
        except Error as e:
            print(f"Error saving comparison to database: {e}")
            connection.rollback()
        finally:
            if connection.is_connected():
                cursor.close()
                connection.close()

def compare_files(file_paths, user_id):
    """Compare files for plagiarism and return results, ignoring the first page of each file."""
    file_contents = [convert_file(file_path, skip_first_page=True) for file_path in file_paths]   
    vectorizer = TfidfVectorizer()
    tfidf_matrix = vectorizer.fit_transform(file_contents)   
    cosine_sim = cosine_similarity(tfidf_matrix)
    results = []
    file_ids = []
    for file_path, content in zip(file_paths, file_contents):
        file_id = save_file_to_db(os.path.basename(file_path), content, user_id)
        if file_id is not None:
            file_ids.append(file_id)
        else:
            print(f"Error: Could not save file {file_path} to database")
    for i in range(len(file_ids)):
        for j in range(i + 1, len(file_ids)):
            similarity = float(cosine_sim[i][j] * 100)
            file1, file2 = file_paths[i], file_paths[j]
            highlighted_file1 = highlight_similar_parts(file_contents[i], file_contents[j])
            highlighted_file2 = highlight_similar_parts(file_contents[j], file_contents[i])
            save_comparison_to_db(file_ids[i], file_ids[j], similarity, user_id, highlighted_file1, highlighted_file2)
            results.append((file1, file2, similarity, highlighted_file1, highlighted_file2))

    return results

def highlight_similar_parts(text1, text2):
    """Highlight similar parts of two texts."""
    tokens1 = text1.split()
    tokens2 = set(text2.split())
    highlighted_tokens = []
    for token in tokens1:
        if token in tokens2:
            highlighted_tokens.append(f'<span style="background-color: yellow;">{token}</span>')
        else:
            highlighted_tokens.append(token)
    return ' '.join(highlighted_tokens)

def get_file_icon(filename):
    extension = filename.split('.')[-1].lower()
    icon_map = {
        'pdf': 'fas fa-file-pdf',
        'docx': 'fas fa-file-word',
        'doc': 'fas fa-file-word',
        'txt': 'fas fa-file-alt',
        'html': 'fas fa-file-code',
    }
    return icon_map.get(extension, 'fas fa-file')

def admin_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if not current_user.is_authenticated or not current_user.is_admin:
            flash('You need to be an admin to access this page.', 'danger')
            return redirect(url_for('index'))
        return f(*args, **kwargs)
    return decorated_function

@app.route('/admin')
@login_required
@admin_required
def admin_dashboard():
    connection = create_database_connection()
    users = []
    comparisons = []
    if connection:
        try:
            cursor = connection.cursor(dictionary=True)
            cursor.execute("SELECT id, username, email, is_admin FROM users")
            users = cursor.fetchall()
            
            cursor.execute("""
                SELECT c.id, f1.filename AS file1, f2.filename AS file2, c.similarity, c.comparison_date, u.username AS user
                FROM comparisons c
                JOIN files f1 ON c.file1_id = f1.id
                JOIN files f2 ON c.file2_id = f2.id
                JOIN users u ON c.user_id = u.id
                ORDER BY c.comparison_date DESC
                LIMIT 10
            """)
            comparisons = cursor.fetchall()
        except Error as e:
            print(f"Error fetching admin dashboard data: {e}")
        finally:
            if connection.is_connected():
                cursor.close()
                connection.close()   
    return render_template('admin_dashboard.html', users=users, comparisons=comparisons)

@app.route('/dashboard')
@login_required
def dashboard():
    connection = create_database_connection()
    if connection:
        try:
            cursor = connection.cursor(dictionary=True)           
            cursor.execute("SELECT COUNT(*) as total FROM comparisons WHERE user_id = %s", (current_user.id,))
            total_comparisons = cursor.fetchone()['total']
            
            cursor.execute("SELECT AVG(similarity) as avg_similarity FROM comparisons WHERE user_id = %s", (current_user.id,))
            avg_similarity_result = cursor.fetchone()['avg_similarity']
            avg_similarity = float(avg_similarity_result) if avg_similarity_result is not None else 0.0
            
            cursor.execute("""
                SELECT 
                    CASE 
                        WHEN similarity <= 30 THEN 'Low'
                        WHEN similarity <= 50 THEN 'Medium'
                        WHEN similarity <= 70 THEN 'High'
                        ELSE 'Vety High'
                    END as category,
                    COUNT(*) as count
                FROM comparisons
                WHERE user_id = %s
                GROUP BY category
            """, (current_user.id,))
            similarity_distribution = {row['category']: row['count'] for row in cursor.fetchall()}           
            cursor.execute("""
                SELECT c.id, f1.filename AS file1, f2.filename AS file2, c.similarity, c.comparison_date
                FROM comparisons c
                JOIN files f1 ON c.file1_id = f1.id
                JOIN files f2 ON c.file2_id = f2.id
                WHERE c.user_id = %s
                ORDER BY c.comparison_date DESC
                LIMIT 10
            """, (current_user.id,))
            recent_comparisons = cursor.fetchall()
            
            return render_template('dashboard.html', 
                                   total_comparisons=total_comparisons,
                                   avg_similarity=avg_similarity,
                                   similarity_distribution=json.dumps(similarity_distribution),
                                   recent_comparisons=recent_comparisons)
        except Error as e:
            print(f"Error fetching dashboard data: {e}")
        finally:
            if connection.is_connected():
                cursor.close()
                connection.close()   
    return render_template('dashboard.html', 
                           total_comparisons=0,
                           avg_similarity=0.0,
                           similarity_distribution=json.dumps({}),
                           recent_comparisons=[])

@app.route('/register', methods=['GET', 'POST'])
def register():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        email = request.form['email']
        # Hash the password using bcrypt
        hashed_password = bcrypt.generate_password_hash(password).decode('utf-8')
        connection = create_database_connection()
        if connection:
            try:
                cursor = connection.cursor()
                cursor.execute("SELECT * FROM users WHERE username = %s", (username,))
                existing_user = cursor.fetchone()
                if existing_user:
                    flash('Username already exists','danger')
                    return redirect(url_for('register'))
                # Insert the new user into the database
                cursor.execute("INSERT INTO users (username, email, password, is_admin) VALUES (%s, %s, %s, %s)",
                               (username, email, hashed_password, False))
                connection.commit()
                flash('Account created successfully! You can now log in.', 'success')
                return redirect(url_for('login'))

            except Error as e:
                print(f"Error during registration: {e}")
                connection.rollback()
                flash('Error during registration. Please try again.', 'danger')
            finally:
                if connection.is_connected():
                    cursor.close()
                    connection.close()

    return render_template('register.html')


@app.route('/api/register', methods=['POST'])
@swag_from('static/docs/register.yml')
def api_register():
    try:
        data = request.get_json()
        username = data.get('username')
        password = data.get('password')
        email = data.get('email')

        if not username or not password or not email:
            return jsonify({"error": "Missing required fields"}), 400

        # Hash the password using bcrypt
        hashed_password = bcrypt.generate_password_hash(password).decode('utf-8')
        connection = create_database_connection()

        if connection:
            cursor = connection.cursor()
            cursor.execute("SELECT * FROM users WHERE username = %s", (username,))
            existing_user = cursor.fetchone()
            if existing_user:
                return jsonify({"error": "Username already exists"}), 409

            # Insert the new user into the database
            cursor.execute("INSERT INTO users (username, email, password, is_admin) VALUES (%s, %s, %s, %s)",
                           (username, email, hashed_password, False))
            connection.commit()

            return jsonify({"message": "User registered successfully"}), 201

    except Exception as e:
        return jsonify({"error": "Error during registration"}), 500

    finally:
        if connection and connection.is_connected():
            cursor.close()
            connection.close()

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        connection = create_database_connection()
        if connection:
            try:
                cursor = connection.cursor(dictionary=True)
                cursor.execute("SELECT * FROM users WHERE username = %s", (username,))
                user = cursor.fetchone()

                if user and bcrypt.check_password_hash(user['password'], password):
                    user_obj = User(id=user['id'], username=user['username'], password=user['password'], is_admin=user['is_admin'])
                    login_user(user_obj)
                    flash('Logged in successfully')
                    if user['is_admin']:
                        return redirect(url_for('admin_dashboard'))
                    else:
                        return redirect(url_for('index'))
                else:
                    flash('Invalid username or password')

            except Error as e:
                print(f"Error during login: {e}")

            finally:
                if connection.is_connected():
                    cursor.close()
                    connection.close()

    return render_template('login.html')

@app.route('/api/login', methods=['POST'])
@swag_from('static/docs/login.yml')
def api_login():
    try:
        # Parse JSON payload
        data = request.get_json()
        if not data:
            return jsonify({"error": "Invalid input, JSON payload required"}), 400
        
        # Extract username and password
        username = data.get('username')
        password = data.get('password')

        # Validate input
        if not username or not password:
            return jsonify({"error": "Username and password are required"}), 400

        # Connect to the database
        connection = create_database_connection()
        if not connection:
            return jsonify({"error": "Database connection failed"}), 500
        
        cursor = connection.cursor(dictionary=True)
        cursor.execute("SELECT * FROM users WHERE username = %s", (username,))
        user = cursor.fetchone()

        # Check if user exists and password is correct
        if user and bcrypt.check_password_hash(user['password'], password):
            # Construct user object for session handling
            user_obj = User(id=user['id'], username=user['username'], password=user['password'], is_admin=user['is_admin'])
            login_user(user_obj)
            return jsonify({"message": "Logged in successfully"}), 200
        else:
            return jsonify({"error": "Invalid username or password"}), 401

    except Exception as e:
        # Log the error for debugging (use proper logging in production)
        print(f"Error during API login: {e}")
        return jsonify({"error": "An unexpected error occurred"}), 500

    finally:
        # Ensure the database connection is closed
        if connection and connection.is_connected():
            cursor.close()
            connection.close()


@app.route('/logout')
def logout():
    logout_user()  
    flash('You have been logged out successfully.', 'success') 
    return redirect(url_for('dashboard'))

@app.route('/api/logout', methods=['POST'])
@swag_from('static/docs/logout.yml')  # Swagger YAML file for logout API
def api_logout():
    try:
        # Log the user out
        logout_user()
        
        # Return success message
        return jsonify({"message": "You have been logged out successfully."}), 200

    except Exception as e:
        print(f"Error during logout: {e}")
        return jsonify({"error": "Error during logout"}), 500

@app.route('/', methods=['GET', 'POST'])
@login_required
def index():
    ensure_upload_folder_exists()

    if request.method == 'POST':
        num_files = int(request.form['num_files'])
        file_paths = []
        for i in range(1, num_files + 1):
            file = request.files.get(f'file_{i}')
            if file:
                file_path = os.path.join(UPLOAD_FOLDER, file.filename)
                file.save(file_path)
                file_paths.append(file_path)

        results = compare_files(file_paths, current_user.id)
        return render_template('results.html', results=results, get_file_icon=get_file_icon)
    
    return render_template('index.html')

@app.route('/api/check-plagiarism', methods=['POST'])
@swag_from('static/docs/check.yml')
def check_plagiarism():
    try:
        # Ensure the upload folder exists
        ensure_upload_folder_exists()

        # Check if 'files' is in the request
        if 'files' not in request.files:
            return jsonify({"error": "No files provided"}), 400

        # Retrieve the files from the form data
        files = request.files.getlist('files')
        user_id = request.form.get('user_id')

        # Validate user_id
        if not user_id or not user_id.isdigit():
            return jsonify({"error": "Invalid or missing user_id"}), 400

        user_id = int(user_id)

        # Ensure the temp_uploads folder exists
        temp_upload_folder = 'temp_uploads'
        if not os.path.exists(temp_upload_folder):
            os.makedirs(temp_upload_folder)

        # Save uploaded files temporarily
        file_paths = []
        for file in files:
            if file.filename == '':
                return jsonify({"error": "One or more files are missing a filename"}), 400

            file_path = os.path.join(temp_upload_folder, file.filename)
            file.save(file_path)
            file_paths.append(file_path)

        # Call the compare_files function
        results = compare_files(file_paths, user_id)

        # Format the response
        response_data = []
        for result in results:
            file1, file2, similarity, highlighted_file1, highlighted_file2 = result
            response_data.append({
                "file1": file1,
                "file2": file2,
                "similarity": similarity,
                "highlighted_file1": highlighted_file1,
                "highlighted_file2": highlighted_file2
            })

        # Clean up temporary files
        for path in file_paths:
            if os.path.exists(path):
                os.remove(path)

        return jsonify(response_data), 200

    except Exception as e:
        print(f"Error during file comparison: {e}")
        return jsonify({"error": "An error occurred during file comparison"}), 500

@app.route('/history')
@login_required
def history():
    page = request.args.get('page', 1, type=int)
    offset = (page - 1) * RECORDS_PER_PAGE

    connection = create_database_connection()
    history_data = []
    total_records = 0
    if connection:
        try:
            cursor = connection.cursor(dictionary=True)
            
            if current_user.is_admin:
                count_query = "SELECT COUNT(*) AS total FROM comparisons"
                cursor.execute(count_query)
                total_records = cursor.fetchone()['total']
                
                data_query = """
                SELECT c.id, f1.filename AS file1, f2.filename AS file2, c.similarity, c.comparison_date,
                       f1.upload_date AS file1_upload_date, f2.upload_date AS file2_upload_date,
                       c.highlighted_content1, c.highlighted_content2,
                       u.username AS user
                FROM comparisons c
                JOIN files f1 ON c.file1_id = f1.id
                JOIN files f2 ON c.file2_id = f2.id
                JOIN users u ON c.user_id = u.id
                ORDER BY c.comparison_date DESC
                LIMIT %s OFFSET %s
                """
                cursor.execute(data_query, (RECORDS_PER_PAGE, offset))
            else:
                count_query = "SELECT COUNT(*) AS total FROM comparisons WHERE user_id = %s"
                cursor.execute(count_query, (current_user.id,))
                total_records = cursor.fetchone()['total']
                
                data_query = """
                SELECT c.id, f1.filename AS file1, f2.filename AS file2, c.similarity, c.comparison_date,
                       f1.upload_date AS file1_upload_date, f2.upload_date AS file2_upload_date,
                       c.highlighted_content1, c.highlighted_content2
                FROM comparisons c
                JOIN files f1 ON c.file1_id = f1.id
                JOIN files f2 ON c.file2_id = f2.id
                WHERE c.user_id = %s
                ORDER BY c.comparison_date DESC
                LIMIT %s OFFSET %s
                """
                cursor.execute(data_query, (current_user.id, RECORDS_PER_PAGE, offset))
            
            history_data = cursor.fetchall()
        except Error as e:
            print(f"Error fetching history: {e}")
        finally:
            if connection.is_connected():
                cursor.close()
                connection.close()

    total_pages = (total_records + RECORDS_PER_PAGE - 1) // RECORDS_PER_PAGE

    return render_template(
        'history.html', 
        history=history_data, 
        page=page, 
        total_pages=total_pages,
        is_admin=current_user.is_admin
    )

@app.route('/delete/<int:id>', methods=['POST'])
def delete(id):
    connection = create_database_connection()
    if connection:
        try:
            cursor = connection.cursor()
            cursor.execute("DELETE FROM comparisons WHERE id = %s AND user_id = %s", (id, current_user.id))
            connection.commit()
            flash('Record deleted successfully', 'success')
        except Error as e:
            print(f"Error deleting history entry: {e}")
        finally:
            if connection.is_connected():
                cursor.close()
                connection.close()
    return redirect(url_for('history'))

@app.route('/help')
def help_page():
    return render_template('help.html')

@app.route('/contact')
@login_required
def contact():
    return render_template('contact.html')

@app.route('/submit_contact', methods=['POST'])
@login_required
def submit_contact():
    name = request.form['name']
    email = request.form['email']
    message = request.form['message']    
    connection = create_database_connection()  
    if connection:
        try:
            cursor = connection.cursor()
            cursor.execute(
                "INSERT INTO contact_messages (name, email, message) VALUES (%s, %s, %s)",
                (name, email, message)
            )
            connection.commit()
            flash('Your message has been sent successfully!', 'success')
        except Exception as e:
            print(f"Error saving contact message: {e}")
            flash('There was an error sending your message. Please try again.', 'error')
        finally:
            if connection.is_connected():
                cursor.close()
                connection.close()

    return redirect(url_for('contact'))  

@app.route('/edit_user/<int:user_id>', methods=['GET', 'POST'])
@login_required
@admin_required
def edit_user(user_id):
    connection = create_database_connection()
    if connection:
        try:
            cursor = connection.cursor(dictionary=True)
            
            if request.method == 'POST':
                new_username = request.form['username']
                new_email = request.form['email']
                is_admin = 'is_admin' in request.form                
                cursor.execute("""
                    UPDATE users 
                    SET username = %s, email = %s, is_admin = %s 
                    WHERE id = %s
                """, (new_username, new_email, is_admin, user_id))
                connection.commit()
                
                flash('User updated successfully', 'success')
                return redirect(url_for('admin_dashboard'))
            
            else:
                cursor.execute("SELECT * FROM users WHERE id = %s", (user_id,))
                user = cursor.fetchone()
                if user:
                    return render_template('edit_user.html', user=user)
                else:
                    flash('User not found', 'danger')
                    return redirect(url_for('admin_dashboard'))
                
        except Error as e:
            print(f"Error in edit_user: {e}")
            flash('An error occurred while editing the user', 'danger')
        finally:
            if connection.is_connected():
                cursor.close()
                connection.close()
    
    return redirect(url_for('admin_dashboard'))

@app.route('/delete_user/<int:user_id>', methods=['POST'])
@login_required
@admin_required
def delete_user(user_id):
    if current_user.id == user_id:
        flash('You cannot delete your own account', 'danger')
        return redirect(url_for('admin_dashboard'))
    
    connection = create_database_connection()
    if connection:
        try:
            cursor = connection.cursor()
            cursor.execute("DELETE FROM comparisons WHERE user_id = %s", (user_id,))
            cursor.execute("DELETE FROM files WHERE user_id = %s", (user_id,))
            cursor.execute("DELETE FROM users WHERE id = %s", (user_id,))            
            connection.commit()
            flash('User and all associated data deleted successfully', 'success')
        
        except Error as e:
            print(f"Error in delete_user: {e}")
            connection.rollback()
            flash('An error occurred while deleting the user', 'danger')
        finally:
            if connection.is_connected():
                cursor.close()
                connection.close()
    
    return redirect(url_for('admin_dashboard'))

@app.route('/view_comparison/<int:comparison_id>')
@login_required
@admin_required
def view_comparison(comparison_id):
    connection = create_database_connection()
    if connection:
        try:
            cursor = connection.cursor(dictionary=True)
            
            query = """
            SELECT c.id, f1.filename AS file1, f2.filename AS file2, 
                   c.similarity, c.comparison_date, u.username AS user,
                   f1.content AS content1, f2.content AS content2
            FROM comparisons c
            JOIN files f1 ON c.file1_id = f1.id
            JOIN files f2 ON c.file2_id = f2.id
            JOIN users u ON c.user_id = u.id
            WHERE c.id = %s
            """
            cursor.execute(query, (comparison_id,))
            comparison = cursor.fetchone()
            
            if comparison:
                return render_template('view_comparison.html', comparison=comparison)
            else:
                flash('Comparison not found', 'danger')
        
        except Error as e:
            print(f"Error in view_comparison: {e}")
            flash('An error occurred while retrieving the comparison', 'danger')
        finally:
            if connection.is_connected():
                cursor.close()
                connection.close()
    
    return redirect(url_for('admin_dashboard'))

@app.route('/delete_comparison/<int:comparison_id>', methods=['POST'])
@login_required
@admin_required
def delete_comparison(comparison_id):
    connection = create_database_connection()
    if connection:
        try:
            cursor = connection.cursor()
            cursor.execute("DELETE FROM comparisons WHERE id = %s", (comparison_id,))           
            connection.commit()
            flash('Comparison deleted successfully', 'success')
        
        except Error as e:
            print(f"Error in delete_comparison: {e}")
            connection.rollback()
            flash('An error occurred while deleting the comparison', 'danger')
        finally:
            if connection.is_connected():
                cursor.close()
                connection.close()
    return redirect(url_for('admin_dashboard'))

if __name__ == '__main__':
    alter_comparisons_table()
    app.run(debug=True)
